import os
from docx import Document
def copy_formatting(base_doc_path, target_folder):
    base_doc = Document(base_doc_path)

    base_formatting = []
    for paragraph in base_doc.paragraphs:
        base_formatting.append((paragraph.style.name, paragraph.style.font))

    for filename in os.listdir(target_folder):
        if filename.endswith(".docx"):
            target_path = os.path.join(target_folder, filename)
            target_doc = Document(target_path)

            for i, paragraph in enumerate(target_doc.paragraphs):
                if i < len(base_formatting):
                    target_style, target_font = base_formatting[i]
                    paragraph.style = target_style
                    for run in paragraph.runs:
                        run.font = target_font

                target_doc.save(target_path)

if __name__ == "__main__":
    base_document_path = r"C:\need\kchay\ПапкаW\Документ1.docx"
    target_folder_path = r"C:\need\kchay\ПапкаC"
    target_folder = os.path.join(target_folder_path, "Папка")

    copy_formatting(base_document_path, target_folder_path)
