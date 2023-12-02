import openpyxl
from docx import Document
from docx.shared import Pt
def generate_invoice(invoice_number):
    wb = openpyxl.load_workbook(r'C:\need\kchay\ExcelApp\ExcelM.xlsx')

    clients_sheet = wb['Покупці']
    products_sheet = wb['Товари']
    invoices_sheet = wb['Рахунки']
    items_sheet = wb['Пункти']

    doc = Document()

    invoice_data = None
    for row in invoices_sheet.iter_rows(values_only=True):
        if row[1] == invoice_number:
            invoice_data = row
            break

    if invoice_data is None:
        print(f"Рахунок з номером {invoice_number} не знайдено.")
        return

    doc.add_heading(f"Рахунок № {invoice_data[1]}", level=1)
    doc.add_paragraph(f"Дата {invoice_data[2].strftime('%d.%m.%Y')}")

    client_data = None
    for row in clients_sheet.iter_rows(values_only=True):
        if row[0] == invoice_data[3]:
            client_data = row
            break

    if client_data is not None:
        doc.add_paragraph(f"Покупець: {client_data[1]}")

    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'

    table.rows[0].cells[0].text = '№'
    table.rows[0].cells[1].text = 'Назва'
    table.rows[0].cells[2].text = 'Од. виміру'
    table.rows[0].cells[3].text = 'Кількість'
    table.rows[0].cells[4].text = 'Ціна'
    table.rows[0].cells[5].text = 'Сума'

    total_amount = 0
    row_number = 1
    for row in items_sheet.iter_rows(values_only=True):
        if row[0] == invoice_number:
            product_data = None
            for product_row in products_sheet.iter_rows(values_only=True):
                if product_row[0] == row[1]:
                    product_data = product_row
                    break

            if product_data is not None:
                quantity = row[2]
                price = product_data[3]
                total = quantity * price

                total_amount += total

                row_cells = table.add_row().cells
                row_cells[0].text = str(row_number)
                row_cells[1].text = product_data[1]
                row_cells[2].text = product_data[2]
                row_cells[3].text = str(quantity)
                row_cells[4].text = f"{price:.2f}"
                row_cells[5].text = f"{total:.2f}"

                row_number += 1

        doc.add_paragraph(f"Всього {total_amount:.2f}")

    doc.save(f"Invoice_{invoice_number}.docx")

if __name__ == "__main__":
    invoice_number_input = input("Введіть номер рахунку: ")

    wb = openpyxl.load_workbook(r'C:\need\kchay\ExcelApp\ExcelM.xlsx')
    invoices_sheet = wb['Рахунки']

    invoice_exists = any(str(row[1]) == invoice_number_input for row in invoices_sheet.iter_rows(values_only=True))

    if not invoice_exists:
        print(f"Рахунок з номером {invoice_number_input} не знайдено.")
    else:
        generate_invoice(invoice_number_input)











